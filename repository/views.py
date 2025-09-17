from django.shortcuts import render, redirect,  get_object_or_404
from django.contrib.auth import authenticate, login, logout
from django.contrib import messages
from django.views.decorators.csrf import csrf_exempt
from .forms import *
from collections import defaultdict
from django.http import JsonResponse, HttpResponseBadRequest
from django.contrib.auth import get_user_model
from django.core.mail import send_mail
from django.utils.timezone import localtime, now
import random, string
from django.db.models import Q
from django.conf import settings
import os
from django.http import HttpResponse, JsonResponse
from .models import *
from django.utils import timezone
import shutil
from datetime import datetime, timedelta
from django.template.loader import render_to_string
from django.contrib.auth.models import User, Group
from django.contrib.auth.hashers import make_password, check_password
from django.contrib.auth.decorators import login_required
# Create your views here.
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
import json, time
import torch
import re
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from .utils import get_nav_context
# Compute similarity
from sklearn.feature_extraction.text import TfidfVectorizer

from sklearn.metrics.pairwise import cosine_similarity

# from torch.nn.functional import cosine_similarity

from .forms import UserForm

from transformers import pipeline
from sentence_transformers import SentenceTransformer
from transformers import pipeline
import pdfplumber
import docx

summarizer = pipeline("summarization", model="facebook/bart-large-cnn")

def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file"""
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(docx_path):
    """Extract text from a Word (.docx) file"""
    doc = docx.Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])

def summarize_text(text, max_length=150, min_length=50):
    """Summarize text using a transformer model"""
    if len(text.split()) < 50:  # If the text is too short, return as is
        return text
    
    try:
        summary = summarizer(text[:1024], max_length=max_length, min_length=min_length, do_sample=False)
        if summary and isinstance(summary, list) and "summary_text" in summary[0]:  
            return summary[0]["summary_text"]
        else:
            return "Summary could not be generated."
    except Exception as e:
        return f"Error: {str(e)}"


@login_required(login_url='Myhome')
def add_user(request):
    user_groups = request.user.groups.values_list('name', flat=True)
    users = User.objects.all().values('id','first_name', 'last_name', 'username', 'password', 'groups__name')
    users_list = list(users)
    
    if request.method == 'POST':
        form = UserForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'User has been added successfully.')
            return render(request, 'add_user.html', {'form': form})
            # return redirect('index')  # Redirect to a user list or success page
    else:
        form = UserForm()
        messages.error(request, 'Please correct the errors below.')
    context = {'users': users_list, 'form':form, 'user_groups':user_groups}
    return render(request, 'add_user.html', context)

def deleteUser(request, id):
    deltodoID = get_object_or_404(User, id=id)

    # Fetch the username from the user object
    fetchUserDetails = deltodoID.username

    # Delete associated UserDetails records
    UserDetails.objects.filter(myuser__username=fetchUserDetails).delete()

    # Delete the user object
    deltodoID.delete()

    
    return redirect('add_user')


@login_required(login_url='Myhome')
def dashboard(request):
    user = request.user
    first_name = user.first_name.strip()  # Get first name
    last_name = user.last_name.strip()    # Get last name
    full_name = f"{first_name} {last_name}"  # Combine first and last name

    
    all_files = uploadFile.objects.all()
    user = request.user
    user_groups = user.groups.values_list('name', flat=True)
    full_name = f"{user.first_name} {user.last_name}" 

    if user.groups.filter(name="Admin").exists() or user.groups.filter(name="SuperAdmin").exists():
        # Admin and SuperAdmin see all files
        files = all_files
    elif user.groups.filter(name="Lecturer").exists():
        full_name = f"{user.first_name} {user.last_name}".strip()
        files = all_files.filter(lecturer=full_name)
    elif user.groups.filter(name="Student").exists():
        files = all_files.filter(nameUploaded=user) #Assuming `uploaded_by` is a ForeignKey to User
    user_groups = request.user.groups.values_list('name', flat=True)
    
    nav_context = get_nav_context(request)
    now = timezone.now()
    meetings = Event_schedule.objects.filter(event_date__gte=now)

    meetings_data = []
    for meeting in meetings:
        # Handle attendees as a JSON string
        attendees_str = meeting.attendees
        try:
            # Convert JSON string to a Python list
            attendees_list = json.loads(attendees_str)
        except json.JSONDecodeError:
            # Handle invalid JSON
            attendees_list = []

        # Count the number of attendees
        attendee_count = len(attendees_list)

        # Collect data for each meeting
        meetings_data.append({
            'meeting': meeting,
            'attendee_count': attendee_count,
            'attendees_list': attendees_list,
        })
    # Pass the meeting and attendee count to the template
    
    docx = DocRepo.objects.all().count()
    fold = FolderRepo.objects.all().count()
    
    user_groups = request.user.groups.values_list('name', flat=True)
    countHour = range(00, 13)
    countMinute = range(00, 60)
    doc = FolderRepo.objects.all().distinct()
    attendee_count = []
     
    
    context = {**nav_context, 'firstname':first_name,'allfiles':files,'user_groups':user_groups,'attendee_count': attendee_count, 'meetings_data': meetings_data,'user_groups':user_groups, 'titeHour':countHour, 'titeMin':countMinute,'fild':doc, 'docx':docx, 'fold':fold}
    return render(request, 'index.html', context)


    

def Myhome(request):
    if request.method == "POST":
        form = LoginForm(request.POST)

        if form.is_valid():
            email = form.cleaned_data['username']
            checkUname = User.objects.filter(username=email)
            if checkUname.exists():
                user_obj = checkUname.first()  # Retrieve the first (and presumably only) result
                username = user_obj.username
                password = form.cleaned_data['password']
                user = authenticate(request, username=username, password=password)

                if user is not None:
                    login(request, user)
                    return redirect("dashboard")
                else:
                    messages.error(request, "Invalid Username or Password")
            else:
                messages.error(request, 'Email not found')
    return render(request, 'login.html')

#logout


def logout(request):
    # Set the timeout duration (in seconds)
    timeout_seconds = 100  # Adjust the timeout to your desired value

    # Get the last activity time from the session or set it to now if not present
    last_active_time = request.session.get('last_activity_time', time.time())

    # If the user has been inactive for the specified timeout period, log them out
    if time.time() - last_active_time > timeout_seconds:
        logout(request)
        # Redirect to the last visited page or a default page
        return redirect(request.META.get('HTTP_REFERER', 'Myhome'))  # Replace 'Myhome' with your default URL

    # Update the session's last activity time
    request.session['last_activity_time'] = time.time()
      # Return None to continue processing the request
    
    # If the user is still active, render the login page
    return redirect('Myhome')

def yes(request):
    if request.method =="POST":
        projectN = request.POST['projectName']
        mycreate = FolderRepo.objects.filter(foldername=projectN)
        if mycreate:
            return redirect('uploadproject')
        else:
            dateY = datetime.now()
            create = FolderRepo(foldername=projectN, date_created=dateY)
            create.save()

    # Define the specific location where you want to create the directory
        directory_path = os.path.join(settings.STATICFILES_DIRS[0], 'medias', projectN)
        
        # Check if the directory exists at the specified location
        if os.path.isdir(directory_path):
             return redirect('uploadproject')
        else:
            # Create the directory if it doesn't exist
            os.makedirs(directory_path)
            # Add a success message to be displayed in the template
            messages.success(request, f"Directory {projectN} created successfully at {directory_path}.")
    
    # Render the template with the success message
    return render(request, "yes.html")
#repository\static

def generate_otp(length=6):
    digits = string.digits
    otp = ''.join(random.choice(digits) for _ in range(length))
    return otp



    

def  forgetpassword(request):
    if request.method == "POST":
        mainemail = request.POST['email']
        request.session['resend_code'] = mainemail
        
        

        user = request.user
        
        base_uri= request.build_absolute_uri()
        userEmail = User.objects.filter(email=mainemail)
        if userEmail:
            otp = generate_otp()
            dateY = datetime.now()
            
            checkMail = PasswordResetCode.objects.update_or_create(email=mainemail,defaults = {
                'otpCode':otp,
                'date_created':dateY
                
            })
            if checkMail:
                
                send_mail("The context form subject", f"Your OTP for reset password is {otp}", "wonderpaul242@gmail.com", ['wonderpaul243@gmail.com'])
                            
                messages.success(request, 'Email sent successfully!')
                return redirect('buttonClickreset')
             
               
                
        else:
            messages.error(request, 'Oops email not found')
            return render(request, "forgetpassword.html")
        # if userEmail:
            
        #     try:
        #         otp = generate_otp()
        #         dateY = datetime.now()
        #         create = PasswordResetCode(email=mainemail, otpCode=otp, date_created=dateY)
        #         ch = PasswordResetCode.objects.get(email=mainemail)
        #         if create.DoesNotExist():
        #             create.save()
                

        #         if create is not None:
                    
                
        #             checkMail = PasswordResetCode.objects.get(email=mainemail)
        #             if checkMail is not None:
        #                 checkMail.email = str(mainemail)
        #                 checkMail.otpCode = otp
        #                 dateY = datetime.now()
        #                 checkMail.date_created = dateY
        #                 checkMail.save()
        #                 # create = PasswordResetCode(email=email, otpCode=otp, date_created=dateY)
        #                 # create.save()
        #                 html = render_to_string("yes.html", {
        #                     "email":mainemail,
        #                     "otp":otp
        #                 })
        #                 try:
                            
        #                     send_mail("The context form subject", f"Your OTP for reset password is {otp}", "wonderpaul242@gmail.com", ['wonderpaul243@gmail.com'])
                            
        #                     messages.success(request, 'Email sent successfully!')
        #                     return redirect('buttonClickreset')
        #                 except Exception as e:
        #                     messages.error(request, f'Failed to send email: {e}')
        #                 return redirect('buttonClickreset')
        #             else:
        #                 messages.error(request, "Email not found")
        #         else:
        #             create.save()
        #     except User.DoesNotExist:
        #         return render(request, "forgetpassword.html")

            
        
        # else:
        #     messages.error(request, "Error in Email or not found, Check and Retype")
        #     return render(request, "forgetpassword.html",)




    return render(request, "forgetpassword.html",)


def buttonClickreset(request):
    if request.method == "POST":
        otpcode = request.POST['otpcode']
        if PasswordResetCode.objects.filter(otpCode=otpcode).exists():
        
            return redirect('resetpassword')
        else:
            messages.error(request, 'Oops Code invalid')
            return render(request, 'buttonClickreset.html')
    objotp = PasswordResetCode.objects.values_list().last()
    drmail = request.session.get('resend_code')
    
    request.session.set_expiry(60)

    context = {'lastOtp':objotp, 'drmail':drmail,}
    return render(request, 'buttonClickreset.html', context)


def resetpassword(request):
    drmail = request.session.get('resend_code')
    if 'resend_code' not in request.session:
        messages.error(request, 'Your session has expired. Please log in again.')
        return redirect('forgetpassword')

    if request.method == "POST":
        passwordx = request.POST['password']
        repassword = request.POST['repassword']
        
        if len(passwordx) < 8:
            messages.error(request, 'password too short. Must be atleast 8.')
            return render(request, 'resetpassword.html')
        hashpass = make_password(passwordx)
        if passwordx == repassword:
            update = User.objects.get(email=drmail)
            update.password = hashpass
            update.save()
            return redirect("yes")
                
        else:
            messages.error(request,"Password does not match")
    return render (request, "resetpassword.html", {'drmail':drmail})


def resendcode(request):
    drmail = request.session.get('resend_code')
    otp = generate_otp()
    dateY = datetime.now()
            
    checkMail = PasswordResetCode.objects.update_or_create(email=drmail,defaults = {
        'otpCode':otp,
        'date_created':dateY
                
    })
    if checkMail:
                
        send_mail("The context form subject", f"Your OTP for reset password is {otp}", "wonderpaul242@gmail.com", ['wonderpaul243@gmail.com'])
                            
        messages.success(request, 'Email sent successfully!')
        return redirect('buttonClickreset')

    return render(request, "resendcode.html", {'drmail':drmail})


def create_event(request):
    if request.method == 'POST':
        title = request.POST['title']
        start_time = request.POST['start_time']
        end_time = request.POST['end_time']
        description = request.POST['description']

        if title and start_time and end_time:  # Add any other necessary validation
            create = Event(title="ok", start_time=start_time, end_time=end_time, description=description)
            create.save()
            messages.success(request, 'Event created successfully!')
            return redirect('calendar_view')  # Adjust the redirect URL as needed
        else:
            messages.error(request, 'Please fill in all required fields.')

    return render(request, 'event.html')


def yes(request):
    if request.method =="POST":
        projectN = request.POST['projectName']
        mycreate = FolderRepo.objects.filter(foldername=projectN)
        if mycreate:
             return redirect('uploadproject')
        else:
            dateY = datetime.now()
            create = FolderRepo(foldername=projectN, date_created=dateY)
            create.save()

    # Define the specific location where you want to create the directory
        directory_path = os.path.join(settings.STATICFILES_DIRS[0], 'medias', projectN)
        
        # Check if the directory exists at the specified location
        if os.path.isdir(directory_path):
             return redirect('uploadproject')
        else:
            # Create the directory if it doesn't exist
            os.makedirs(directory_path)
            # Add a success message to be displayed in the template
            # messages.success(request, f"Directory {projectN} created successfully at {directory_path}.")
            messages.success(request, 'Project created successfully!')
            return redirect("uploadproject")
    
    # Render the template with the success message
    return render(request, "uploadproject.html")
#repository\static

def uploadproject(request):
    nav_context = get_nav_context(request)
    dj = FolderRepo.objects.all().distinct()
    user_groups = request.user.groups.values_list('name', flat=True)
    if request.method == 'POST' and request.FILES.get('file'):
        uploaded_file = request.FILES['file']
        
        # Define the custom directory path here
        custom_dir = request.POST.get('custom_dir', 'default')
        name_of_project = request.POST.get('name_of_project', 'default')
        deupload_path = os.path.join(settings.STATICFILES_DIRS[0], f'medias/{custom_dir}')
        upload_path = os.path.join(settings.STATICFILES_DIRS[0], f'medias/{custom_dir}', name_of_project)
        
        # Ensure the directory exists
        if not os.path.exists(upload_path):
            os.makedirs(upload_path)
            messages.error(request, {'status': f"File Exist {custom_dir}!"})
        # Save the file
        file_path = os.path.join(upload_path, uploaded_file.name)
        with open(file_path, 'wb+') as destination:
            for chunk in uploaded_file.chunks():
                destination.write(chunk)
    
        
        # return JsonResponse({'status': f"File uploaded successfully to {deupload_path}!"})
    context = {**nav_context,'directory':dj, 'user_groups':user_groups}

    return render(request, "uploadproject.html", context)


def updateProfile(request):
    if request.method == "POST":
        myus =  request.POST.get('myuser')
        yh = User.objects.get(username = myus)

        name =  request.POST['name']
        title =  request.POST['title']
        location =  request.POST['location']
        status =  request.POST['status']
        image =  request.FILES['profile']
        
        if image =='':
           ct =  UserDetails.objects.get(myuser=yh)
           image = ct.profile
        else:
            
            create = UserDetails.objects.update_or_create(myuser=yh, 
                                                          defaults=
                                                          {'name':name,
                                                           'title':title, 'location':location, 'status':status, 'profile':image})
            if create is not None:
                print("hh")
            else:
                messages.error(request, "failed")
    return render(request, 'uploadproject.html')

@login_required(login_url='Myhome')
def projectfolder(request):
    if request.method=="POST":
        folderName = request.POST['foldername']
        
        ParentName = request.POST['parentname']
        dateY = datetime.now()
        addfolder = FolderRepo.objects.filter(foldername=folderName).exists()
        # create = FolderNameProject.objects.update_or_create(name=addfolder, 
        #                                                   defaults=
        #                                                   {'foldername':folderName,
        #                                                    'parentName':ParentName,date_created':dateY})
        #
        if addfolder:
            messages.success(request, 'Project Created Successfully')
            return redirect('uploadproject')                                                    
        else:
                
                
                folder_path = os.path.join(settings.STATICFILES_DIRS[0], f'medias/{ParentName}', folderName)
                folder_path_uri = folder_path.replace("%20","")
                create = FolderNameProject(foldername=folderName, parentname=ParentName, folderurl=folder_path_uri, date_created=dateY)
                create.save()
                
                try:
                    idexit = FolderNameProject.objects.filter(foldername=ParentName).exists()
                    
                    if idexit is not None:
                        os.makedirs(folder_path, exist_ok=True)
                    else:
                        
                       
                        # return JsonResponse({'message': f'Folder "{folderName}" created successfully!'})
                        return redirect('uploadproject')
                        # return redirect("projectfolder")
                        
                        # return JsonResponse({'message': 'Invalid request'}, status=400)
                        
                        # return redirect('projectfolder')
                   
                        
                except Exception as e:
                    return JsonResponse({'message': f'Failed to create folder: {str(e)}'}, status=500)
                    create = FolderNameProject(foldername=folderName, parentname=ParentName, date_created=dateY)
                    # if create:
                    #     create.save()
                        
                    #     # return JsonResponse({'message': 'Invalid request'}, status=400)
                        
                    #     return redirect('projectfolder')
                

    return render(request, 'uploadproject.html')

def delete_folder(request, id):
  folder = get_object_or_404(FolderRepo, id=id)
  folder_path =  os.path.join(settings.STATICFILES_DIRS[0], "medias", folder.foldername)
  if os.path.exists(folder_path):
    shutil.rmtree(folder_path)
    folder.delete()
    # return JsonResponse({'message': f'Folder "{folder.foldername}" and its directory deleted successfully!'})
    return redirect('uploadproject')
  else:
    return JsonResponse({'message': 'Directory does not exist'}, status=400)
  

  
def uploadDocProject(request, id):
    nav_context = get_nav_context(request)
    dj = FolderRepo.objects.get(id=id)
    df = dj.id
    request.session['resend_id'] = df;
    request.session['res_id'] = df;
    djx = FolderNameProject.objects.filter(parentname=dj.foldername)
    if djx:
        djy = ""
        for djy in djx:
            pna = djy.parentname
            ffdn = djy.foldername
            fdn = dj.foldername
            fdname = FolderNameProject.objects.filter(parentname=fdn)
            myfile = uploadFile.objects.filter(parentname=pna, foldername=ffdn)
            context = {'directory':dj, 'fdname':fdname, 'myfile':myfile}
            return render(request, 'uploadDocProject.html', context)
    else:
        from django.utils.safestring import mark_safe
        messages.success(request, mark_safe("No Record Added for this Specific Folder <a href='/repo/uploadproject/'> << Go back and upload</a>"))
    # if djx:
    #     djy = FolderNameProject.objects.get(parentname=dj.foldername)
    
    #     pna = djy.parentname
    #     ffdn = djy.foldername
    #     fdn = dj.foldername
    #     fdname = FolderNameProject.objects.filter(parentname=fdn)
    #     myfile = uploadFile.objects.filter(parentname=pna, foldername=ffdn)
    #     user_groups = request.user.groups.values_list('name', flat=True)
    #     context = {'directory':dj, 'fdname':fdname, 'myfile':myfile}
    # else:
    #     messages.error(request, "No Stuff")
    
        return render(request, 'uploadDocProject.html')
    



def renamefolder(request):
    nav_context = get_nav_context
    if request.method == "POST":
        idfolder = request.POST.get('folderId')
        foldername = request.POST['foldername']
        oldfoldername = request.POST['oldfoldername']
        getfolder = FolderNameProject.objects.get(id=idfolder)
        if getfolder:
            oldfolder_path =  os.path.join(settings.STATICFILES_DIRS[0], f"medias/{getfolder.parentname}", oldfoldername)
            newfolder_path =  os.path.join(settings.STATICFILES_DIRS[0], f"medias/{getfolder.parentname}", foldername)
            getfolderx = uploadFile.objects.filter(foldername__icontains=oldfoldername)
            for x in getfolderx:
                
                x.foldername = x.foldername.replace(oldfoldername, foldername)
                x.save()
            if os.path.exists(oldfolder_path):
                os.rename(oldfolder_path, newfolder_path)
            

            getfolder.foldername = foldername
            getfolder.save()
            id = request.session.get('resend_id')
    context = {**nav_context}
    return redirect(f'../geteachDio/{id}', context)

def deletefolder(request, id):
    folder = FolderNameProject.objects.get(id=id)
    id = request.session.get('resend_id')
    idf = request.session.get('res_id')
    parentname = folder.parentname
    foldername = folder.foldername
    myfolder =  foldername.strip()
    folder_path =  os.path.join(os.path.join(settings.STATICFILES_DIRS[0], f'medias/{parentname}/', myfolder))
    if os.path.exists(folder_path):
        shutil.rmtree(folder_path)
        folder.delete()
        return redirect(f'../uploadDocProject/{idf}')
        # return JsonResponse({'message': f'Folder "{folder.foldername}" and its directory deleted successfully!'})
    else:
        return HttpResponse({'Message':'error'})
      
  
  

def geteachDio(request, id):
    nav_context = get_nav_context(request)
    user_groups = request.user.groups.values_list('name', flat=True)
    dj = FolderNameProject.objects.get(id=id)
    df = dj.foldername
    dc = dj.id
    request.session['resend_id'] = dc
    djx = uploadFile.objects.filter(foldername=df)
    dfuri = uploadFile.objects.filter(foldername=df).last()
    help = "help"
    
    context = {**nav_context,'dj':dj, 'help':help, 'djx':djx, 'dfuri':dfuri,'df':df, 'dc':dc, 'user_groups':user_groups}
      
        

    # if djx:
    #     djy = FolderNameProject.objects.get(parentname=dj.foldername)
    
    #     pna = djy.parentname
    #     ffdn = djy.foldername
    #     fdn = dj.foldername
    #     fdname = FolderNameProject.objects.filter(parentname=fdn)
    #     myfile = uploadFile.objects.filter(parentname=pna, foldername=ffdn)
    #     user_groups = request.user.groups.values_list('name', flat=True)
    #     context = {'directory':dj, 'fdname':fdname, 'myfile':myfile}
    # else:
    #     messages.error(request, "No Stuff")
    
    return render(request, 'geteachDio.html', context)



def document(request):
    nav_context = get_nav_context(request)
    user_groups = request.user.groups.values_list('name', flat=True)
    allDoc = DocRepo.objects.all().distinct()
    context = {**nav_context, 'myDoc':allDoc, 'user_groups':user_groups}
    if request.method =="POST":
        docN = request.POST['documentname']
        mycreate = DocRepo.objects.filter(documentname=docN)
        if mycreate:
            return HttpResponse("Exists")
        else:
            dateY = datetime.now()
            create = DocRepo(documentname=docN, date_created=dateY)
            create.save()

    # Define the specific location where you want to create the directory
        directory_path = os.path.join(settings.STATICFILES_DIRS[0], 'documents', docN)
        
        # Check if the directory exists at the specified location
        if os.path.isdir(directory_path):
            return HttpResponse("Exists")
        else:
            # Create the directory if it doesn't exist
            os.makedirs(directory_path)
            # Add a success message to be displayed in the template
            # messages.success(request, f"Directory {projectN} created successfully at {directory_path}.")
            messages.success(request, 'Document created successfully!')
            return redirect("document")
    return render(request, 'document.html', context)


def message(request, id):
    nav_context = get_nav_context(request)
    user_groups = request.user.groups.values_list('name', flat=True)
    now = datetime.now()
    dateY = now.strftime("%Y-%m-%d %H:%M")
    docfile = DocRepo.objects.get(id=id)
    request.session['mymessageid'] = docfile.id
    docn = docfile.documentname
    messageCk = MessageBox.objects.exclude(SaveToDraft='SaveToDraft').filter(documentName=docn).order_by('messageTitle').distinct()
    if docfile is not None:
        if request.method == "POST":
            docN = request.POST['documentName']
            mesT = request.POST['messageTitle']
            mesTx = request.POST['messageText']
            SaveAsD = request.POST['SaveToDraft']
            mydate = request.POST.get('mydate')
            if mydate == "":
                mydate = datetime.now()
            dateY= datetime.now()
            create =  MessageBox(documentName=docN, messageTitle=mesT, messageText=mesTx, SaveToDraft=SaveAsD, mydate=mydate)
            if MessageBox.objects.filter(documentName=docN, messageTitle=mesT, messageText=mesTx, SaveToDraft=SaveAsD):
                messages.error(request, 'Message Exist')
            else:

                if create:
                    create.save()
                    messages.success(request, 'Message Set Successfully')
                else:
                    messages.error(request, 'Failed to send message')
        
    context = {**nav_context, 'docn':docn, 'docfile':docfile, 'messageCk':messageCk, 'user_groups':user_groups, 'dateT':dateY}    

    return render(request, 'message.html', context)


def viewmessage(request, id):
    nav_context = get_nav_context(request)
    docfile = MessageBox.objects.get(id=id)
    docT = docfile.messageText
    messageCk = MessageBox.objects.filter(messageText=docT).order_by('messageTitle').distinct()
    
    context = {**nav_context, 'messageCk':messageCk}
    return render(request, 'viewmessage.html', context)

def uploadDocument(request):
    return render(request, 'index.html')


def draftmessages(request):
    nav_context = get_nav_context(request)
    user_groups = request.user.groups.values_list('name', flat=True)
    dateY = datetime.now()
    current_time = dateY
    messageCk = MessageBox.objects.exclude(SaveToDraft="Direct Post")
    if request.method == "POST":
        if 'submitformone' in request.POST:
            docN = request.POST['documentName']
            ddid = request.POST['id']
            mesT = request.POST['messageTitle']
            mesTx = request.POST['messageText']
            SaveAsD = request.POST['SaveToDraft']
            mydate = request.POST['mydate']
            if mydate == "":
                mydate = datetime.now()
                dateY= datetime.now()
            ck = MessageBox.objects.get(id=ddid)
            if ck:
                ck.messageText = mesTx
                ck.messageTitle = mesT
                ck.mydate = mydate
                ck.SaveToDraft = SaveAsD
                ck.save()
                messages.success(request, 'Message Set Successfully')
            else:
                messages.error(request, 'Failed to send message')
        

    
    context = {**nav_context,'messageCk':messageCk, 'user_groups':user_groups, 'cur':current_time}
    return render(request, 'draftmessages.html', context)

def deletemessages(request, id):
    mymessageid = request.session.get('mymessageid')
    docfile = MessageBox.objects.get(id=id)
    docfile.delete()
    return redirect(f'message/{mymessageid}')



def deletemymessage(request, id):
    mymessageid = request.session.get('mymessageid')
    docfile = MessageBox.objects.get(id=id)
    docfile.delete()
    return redirect('draftmessages')

def todos(request, id):
    nav_context = get_nav_context(request)
    user_groups = request.user.groups.values_list('name', flat=True)
    users = User.objects.all()
    docfile = DocRepo.objects.get(id=id)
    docN = docfile.documentname
    todos = Todos.objects.filter(documentName=docN).distinct()
    if request.method == 'POST':
        if 'create' in request.POST:

            request.session['todoID'] = docfile.id
            tasklist = request.POST['tasklist']
            taskdetails = request.POST['taskdetails']
            assignee = request.POST['assignee']
            dueon = request.POST['dueon']
            if Todos.objects.filter(tasklist=tasklist,taskdetails=taskdetails,assignee=assignee,dueon=dueon,documentName=docN):
                messages.error(request, 'Record Added Already')
                return render(request, 'todos.html')
            else:

                create = Todos(tasklist=tasklist,taskdetails=taskdetails,assignee=assignee,dueon=dueon,documentName=docN)
                if create:
                    todoid = request.session.get('todoID')
                    create.save()
                
                
                else:
                    messages.error(request, 'Todo Failed to add')
        elif 'update_record' in request.POST:
            request.session['todoID'] = docfile.id
            tasklist = request.POST['tasklist']
            taskid = request.POST['id']
            taskdetails = request.POST['taskdetails']
            assignee = request.POST['assignee']
            dueon = request.POST['dueon']
            getRecord = Todos.objects.get(id=taskid)
            if getRecord:
                getRecord.tasklist = tasklist
                getRecord.taskdetails = taskdetails
                getRecord.assignee = assignee
                getRecord.dueon = dueon
                getRecord.save()


    context = {**nav_context, 'docN':docN, 'user_groups':user_groups, 'todos':todos,'users':users}
    return render(request, 'todos.html', context)

def deleteTodo(request, id):
    deltodoID =  Todos.objects.get(id=id)
    deltodoID.delete()
    todo_id=request.session.get('todoID')
    return redirect(f'todos/{todo_id}')


def deleteCalender(request, id):
    deltodoID =  Event_schedule.objects.get(id=id)
    deltodoID.delete()
    
    return redirect('calendar')




def chat(request):
    nav_context = get_nav_context(request)
    user_groups = request.user.groups.values_list('name', flat=True)
    Alluser = User.objects.values()
    datej = datetime.now()
    month = datej.strftime("%B")
    day = datej.strftime("%A")
    dayy = datej.day
    year = datej.year
    chatIn = ChatIncoming.objects.all()
    initial_data = User.objects.all()
    allDate = f'{day}, {month}, {dayy}'
    if request.method == "POST":
        if 'sendChat' in request.POST:
            user = request.POST['user']
            theuser = user.replace('@', '')
            mymessage = request.POST['mymessage']
            users = User.objects.get(username = theuser)
            if users:
                mycheckuser = UserDetails.objects.get(myuser__username = users.username)
                profile = mycheckuser.profile
                name = mycheckuser.name
            dateY  = datetime.now()
            created = ChatIncoming(mymessage=mymessage, name=name, myuser=users, mydate=dateY, profile=profile)
            created.save()
        elif request.headers.get('x-requested-with') == 'XMLHttpRequest':
            search = request.GET.get('username', '')
            users = User.objects.filter(username__icontains = search)
            results = [{'username': user.username, 'email': user.email} for user in users]
            
            return JsonResponse(results, safe=False)
    
    
    # payload = []
    # for user in users:
    #         payload.append({
    #             'username':user.username
    #         })
        # return JsonResponse({
        #     'status':True,
        #     'payload':payload
        # })
    context = {**nav_context, 'allDate':allDate, 'item':initial_data,'user_groups':user_groups, 'chatIn':chatIn, 'theuser':Alluser}
    return render(request, 'chat.html', context)


def mychat(request):
    user_groups = request.user.groups.values_list('name', flat=True)
    Alluser = UserDetails.objects.all()
    x =  None
    for x in Alluser:
        Tusername= x.myuser.username
    datej = datetime.now()
    month = datej.strftime("%B")
    day = datej.strftime("%A")
    dayy = datej.day
    year = datej.year
    chatIn = ChatIncoming.objects.all()
    x =  None
    for x in chatIn:
        usernameX = x.myuser.username
        # if usernameX = 
    initial_data = User.objects.all()
    allDate = f'{day}, {month}, {dayy}'
    context = {'allDate':allDate, 'item':initial_data,'user_groups':user_groups, 'chatIn':chatIn, 'theuser':Alluser, 'username':Tusername}
    return render(request, 'mychat.html', context)




@login_required(login_url='Myhome')
def chatroom(request):
    nav_context = get_nav_context(request)
    user_groups = request.user.groups.values_list('name', flat=True)
    if request.method == 'POST':
        receiver_id = request.POST.get('receiver_id')
        content = request.POST.get('content')
        uploaded_file = request.FILES.get('file')

        receiver = get_object_or_404(User, id=receiver_id)

        Message.objects.create(
            sender=request.user,
            receiver=receiver,
            content=content,
            file=uploaded_file
        )

        return JsonResponse({'success': True})
    
    users = UserDetails.objects.exclude(myuser__username=request.user.username)
    context = {**nav_context, 'users': users, 'user_groups':user_groups}
    return render(request, 'chatroom.html', context)

@login_required(login_url='Myhome')
def get_chat_history(request):
    if not request.user.is_authenticated:
        return redirect('login')

    user_id = request.GET.get('user_id')
    
    if not user_id:
        return HttpResponseBadRequest('Missing user_id')

    try:
        user = get_object_or_404(User, id=int(user_id))
    except ValueError:
        return HttpResponseBadRequest('Invalid user_id')

    messages = Message.objects.filter(
        (models.Q(sender=request.user) & models.Q(receiver=user)) |
        (models.Q(sender=user) & models.Q(receiver=request.user))
    ).order_by('timestamp')

    chat_history = []
    for msg in messages:
        sender_profile_image = get_object_or_404(UserDetails, myuser=msg.sender).profile.url
        chat_history.append({
            'sender': msg.sender.username,
            'sender_profile': sender_profile_image,
            'content': msg.content,
            'file': msg.file.url if msg.file else None,
            'timestamp': msg.timestamp.strftime('%Y-%m-%d %H:%M:%S')
        })

    selected_user_profile = get_object_or_404(UserDetails, myuser=user).profile.url
    selected_user_name = user.username

    return JsonResponse({
        'chat_history': chat_history,
        'selected_user': {
            'username': selected_user_name,
            'profile_image': selected_user_profile
        }
    })
@login_required(login_url='Myhome')
def activities(request):
    nav_context = get_nav_context(request)
    user_groups = request.user.groups.values_list('name', flat=True)
    files = uploadFile.objects.order_by('-date_created')[:1]
    filesM = Meeting_Schedule.objects.order_by('-date_created')[:1]

    files_with_datesM = []
    for fileM in filesM:
        # theId = FolderNameProject.objects.get(foldername=file.foldername)
        # myid = theId.id
        first_letterM = str(fileM.posted)
        firstLM = first_letterM[0]
        # Extract month and day from the date_created field
        month = fileM.date_created.strftime("%b")  # Full month name (e.g., "August")
        day = fileM.date_created.strftime("%d") 
        time = fileM.date_created.strftime("%I:%M %p")   # Day with leading zero (e.g., "30")

        # Append to the list as a dictionary
        files_with_datesM.append({
            'fileM': fileM,
            'firstlM':firstLM.upper(),
            # 'myid': myid,
            'month': month,
            'day': day,
            'time' : time
        })
    # Create a list to store the month and day for each file
    files_with_dates = []
    theId = []
    file = []
    for file in files:
        theId = FolderNameProject.objects.get(foldername=file.foldername)
        myid = theId.id
        first_letter = str(file.nameUploaded)
        firstL = first_letter[0]
        # Extract month and day from the date_created field
        month = file.date_created.strftime("%b")  # Full month name (e.g., "August")
        day = file.date_created.strftime("%d") 
        time = file.date_created.strftime("%I:%M %p")   # Day with leading zero (e.g., "30")

        # Append to the list as a dictionary
        files_with_dates.append({
            'file': file,
            'firstl':firstL.upper(),
            'myid': myid,
            'month': month,
            'day': day,
            'time' : time
        })
    context = {**nav_context,'getFile':files_with_dates, 'getFileM':files_with_datesM, 'user_groups':user_groups}
    return render(request, 'activities.html', context)



@login_required(login_url='Myhome')
def schedule_meeting(request):
    nav_context = get_nav_context(request)
    myuser = request.user
    user_groups = request.user.groups.values_list('name', flat=True)
    usernames = User.objects.values_list('username', flat=True)

# Filter UserDetails based on the list of usernames
    user_details = UserDetails.objects.filter(myuser__username__in=usernames)
    if request.method == 'POST':
        # Capture form data
        title = request.POST.get('title')
        date = request.POST.get('date')
        start_time = request.POST.get('start_time')
        finish_time = request.POST.get('finish_time')
        agenda = request.POST.get('agenda')
        attendees = request.POST.get('attendees', '[]')  # Get attendees as JSON string

        # Parse the JSON string to a Python list
        try:
            attendees_list = json.loads(attendees)
        except json.JSONDecodeError:
            attendees_list = []

        # Create a new Meeting object and save it to the database
        meeting = Meeting_Schedule(
            posted = myuser,
            title=title,
            date=date,
            start_time=start_time,
            finish_time=finish_time,
            agenda=agenda,
            attendees=attendees_list,
            date_created = datetime.now()
        )
        meeting.save()
    
    

        # Redirect to a success page or any other page after saving
         # Replace 'success' with the name of your success URL or view
    context = {**nav_context, 'ald':user_details, 'user_groups':user_groups}
    # Render the form template for GET requests
    return render(request, 'schedule_meeting.html',context)


def day_with_suffix(day):
    day = int(day)  # Convert to integer
    if 11 <= day <= 13:
        suffix = "th"
    else:
        suffix_map = {1: "st", 2: "nd", 3: "rd"}
        suffix = suffix_map.get(day % 10, "th")
    return day, suffix

def calendar(request):
    nav_context = get_nav_context(request)
    filesM = Event_schedule.objects.all().distinct()
    files_with_datesM = []
    for fileM in filesM:
        
        month = fileM.event_date.strftime("%B")  # Full month name (e.g., "Aug")
        day, suffix = day_with_suffix(fileM.event_date.strftime("%d"))
        time = fileM.event_time.strftime("%I:%M %p")   # Day with leading zero (e.g., "30")

        # Append to the list as a dictionary
        files_with_datesM.append({
            'fileM': fileM,
            
            # 'myid': myid,
            'month': month,
            'day': day,
            'suffix':suffix,
            'time' : time
        })
    myuser = request.user
    user_groups = request.user.groups.values_list('name', flat=True)
    usernames = User.objects.values_list('username', flat=True)

# Filter UserDetails based on the list of usernames
    user_details = UserDetails.objects.filter(myuser__username__in=usernames)
    month = datetime.now().strftime("%B")  # Full month name (e.g., "Aug")
    day, suffix = day_with_suffix(datetime.now().strftime("%d"))

    if request.method == 'POST':
        title = request.POST['title']
        event_date = request.POST['event_date']
        event_time = request.POST['event_time']
        stored_hours = request.POST['stored_hours']
        generated_links = request.POST['generated_links']
        category = request.POST['category']
        attendees = request.POST.get('attendees', '[]') 

        try:
            attendees_list = json.loads(attendees)
        except json.JSONDecodeError:
            attendees_list = []
        myevt = Event_schedule.objects.filter(
            title=title,
            event_date=event_date,
            event_time=event_time,
            stored_hours=stored_hours,
            generated_links=generated_links,
            category=category, 
            attendees = attendees_list,
        )
        if myevt:
            messages.error(request, 'Schedule Already Exist')
            return render(request, 'calendar.html')
        else:
            event = Event_schedule(
                title=title,
                event_date=event_date,
                event_time=event_time,
                stored_hours=stored_hours,
                generated_links=generated_links,
                category=category, 
                attendees = attendees_list,
                date_created  = datetime.now()
                
            )
            event.save()
            return redirect('calendar')
        
        # Add attendees to the event
        
    
    context = {**nav_context, 'month': month, 'day': day, 'suffix': suffix,'getFileM':files_with_datesM, 'ald':user_details, 'user_groups':user_groups}
    return render(request, 'calendar.html', context)


# Load pre-trained model
model = SentenceTransformer('all-MiniLM-L6-v2')

def recommend_projects(user_input):
    if not user_input.strip():  # Handle empty input
        return []

    projects = ProjectIdea.objects.all()
    project_titles = [project.title for project in projects]

    if not project_titles:  # Ensure we have projects to compare
        return []

    # Encode user input & projects
    user_embedding = model.encode(user_input, convert_to_tensor=True)  # Shape: (1, 384)
    project_embeddings = model.encode(project_titles, convert_to_tensor=True)  # Shape: (N, 384)

    # Compute cosine similarity
    scores = cosine_similarity(user_embedding.unsqueeze(0), project_embeddings).squeeze(0)  # Ensure it's 1D

    # Rank and return top 5 projects
    top_indices = torch.argsort(torch.tensor(scores), descending=True)[:5]  # Convert to tensor before sorting
    recommended = [projects[i] for i in top_indices.tolist()]  # Convert indices to project objects

    return recommended

@login_required
def showSuggest(request):
    userGroup = request.user.groups.values_list('name', flat=True)
    context = {'user_groups':userGroup}
    if request.method == "POST":
        user_input = request.POST.get("query", "").strip()
        recommendations = recommend_projects(user_input)
        return render(request, "showSuggest.html", {"recommendations": recommendations})

    return render(request, "showSuggest.html", context)



# Fixed Required Words
FIXED_REQUIRED_WORDS = ["Chapter", "Introduction"]

def clean_text(text):
    """Normalize text by removing extra spaces and newlines"""
    return re.sub(r"\s+", " ", text).strip().lower()

def contains_required_words(text, fixed_words, lecturer_names):
    """Ensure all fixed words and at least one lecturer's name exist in the text"""
    text = text.lower()
    return all(word.lower() in text for word in fixed_words) and any(name.lower() in text for name in lecturer_names)

def calculate_similarity(new_text, existing_texts):
    """Compare new text with existing texts and return max similarity score"""
    if not existing_texts:
        return 0  
    texts = [new_text] + existing_texts
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform(texts)
    similarity_scores = cosine_similarity(tfidf_matrix[0], tfidf_matrix[1:]).flatten()
    return max(similarity_scores) if similarity_scores.size > 0 else 0

def safe_int_conversion(value, default=12):
    """Safely convert a value to an integer"""
    try:
        return int(float(value))  # Converts to float first, then int
    except (ValueError, TypeError):
        return default  # Return default font size if conversion fails

def apply_format_to_pdf(file_path, formatting):
    """Apply document formatting to a PDF file"""
    output_file = file_path.replace(".pdf", "_formatted.pdf")

    c = canvas.Canvas(output_file, pagesize=letter)
    font_size = safe_int_conversion(formatting.fontSize, default=12)  # Ensure font size is valid
    line_spacing = safe_int_conversion(formatting.lineSpacing, default=1.5) * 5  # Convert to int

    c.setFont(formatting.fontFamily, font_size)
    c.setFillColor(colors.HexColor(formatting.fontColor))

    y_position = 750
    with open(file_path, "r", encoding="utf-8") as f:
        for line in f.readlines():
            c.drawString(100, y_position, line.strip())
            y_position -= line_spacing  
            if y_position < 50:
                c.showPage()
                c.setFont(formatting.fontFamily, font_size)
                y_position = 750

    c.save()
    os.replace(output_file, file_path)  # Replace original file

def apply_format_to_docx(file_path, formatting):
    """Apply document formatting to a DOCX file"""
    
    doc = DocxDocument(file_path)

    font_size = safe_int_conversion(formatting.fontSize, default=12)
    letter_spacing = safe_int_conversion(formatting.letterSpacing, default=0)

    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = formatting.fontFamily
            run.font.size = Pt(font_size)
            # run.font.color.rgb = RGBColor.from_string(formatting.fontColor.replace("#", ""))
            # Ensure letter spacing is handled safely
            if hasattr(run.font, "letter_spacing"):
                run.font.letter_spacing = Pt(letter_spacing)

    output_file = file_path.replace(".docx", "_formatted.docx")
    doc.save(output_file)
    os.replace(output_file, file_path)  # Replace original file

@login_required(login_url='Myhome')
def upload_document(request):
    user_groups = request.user.groups.values_list('name', flat=True)
    firstname = request.user.first_name if request.user.is_authenticated else "Guest"

    lecturer_group = Group.objects.filter(name="Lecturer").first()
    lecturers = lecturer_group.user_set.all() if lecturer_group else []
    lecturer_names = [lecturer.get_full_name().lower() for lecturer in lecturers]

    if request.method == "POST":
        form = DocumentForm(request.POST, request.FILES)
        if form.is_valid():
            document = form.save()
            file_path = document.file.path  

            #  Extract text from document
            extracted_text = ""
            if file_path.lower().endswith(".pdf"):
                extracted_text = extract_text_from_pdf(file_path)
            elif file_path.lower().endswith(".docx"):
                extracted_text = extract_text_from_docx(file_path)

            if not extracted_text:
                document.delete()
                messages.error(request, "Could not extract text from the document.")
                return render(request, "upload_document.html", {"form": form})

            extracted_text = clean_text(extracted_text)

            if not contains_required_words(extracted_text, FIXED_REQUIRED_WORDS, lecturer_names):
                document.delete()
                messages.error(request, f"Document must contain all of these words: {', '.join(FIXED_REQUIRED_WORDS)} and at least one lecturer's name.")
                return render(request, "upload_document.html", {"form": form})

            existing_texts = [clean_text(doc.summary) for doc in Document.objects.all() if doc.summary]
            similarity_score = calculate_similarity(extracted_text, existing_texts)

            if similarity_score >= 0.35:  
                document.delete()
                DocumentError = f"A similar project already exists in the database with {similarity_score * 100:.2f}% similarity."
                theSimilarity = f"{similarity_score * 100:.2f}%"
                return render(request, "upload_document.html", {"form": form, 'user_groups': user_groups, 'firstname': firstname, 'DocsError': DocumentError, 'simScore': theSimilarity})

            #  Save summary and upload document
            document.summary = summarize_text(extracted_text)
            document.save()

            #  Apply formatting from formatDocument model
            formatting = formatDocument.objects.first()
            if formatting:
                if file_path.lower().endswith(".pdf"):
                    apply_format_to_pdf(file_path, formatting)
                elif file_path.lower().endswith(".docx"):
                    apply_format_to_docx(file_path, formatting)

            messages.success(request, "Document uploaded and formatted successfully!")
            return redirect("document_list")

    else:
        form = DocumentForm()
    
    return render(request, "upload_document.html", {"form": form, "firstname": firstname})


@login_required(login_url='Myhome')
def upload_file_view(request):
    user_groups = request.user.groups.values_list('name', flat=True)
    firstname = request.user.first_name if request.user.is_authenticated else "Guest"
    id = request.session.get('resend_id')

    # Get lecturer names
    lecturer_group = Group.objects.filter(name="Lecturer").first()
    lecturers = lecturer_group.user_set.all() if lecturer_group else []
    lecturer_names = [lecturer.get_full_name().lower() for lecturer in lecturers]

    if request.method == 'POST':
        try:
            nameUploaded = request.user
            filename = request.POST.get('filename')
            foldername = request.POST.get('foldername')
            parentname = request.POST.get('parentname')
            lecturer = request.POST.get('lecturer')
            status = "Not Approve"
            dateY = timezone.now()

            document = request.FILES.get('image')
            if not document:
                messages.error(request, "No document uploaded.")
                return redirect(f'../geteachDio/{id}')

            # Save uploaded file temporarily
            temp_path = os.path.join(settings.MEDIA_ROOT, 'temp', document.name)
            os.makedirs(os.path.dirname(temp_path), exist_ok=True)
            with open(temp_path, 'wb+') as temp_file:
                for chunk in document.chunks():
                    temp_file.write(chunk)

            #  Extract and clean text
            extracted_text = ""
            if temp_path.lower().endswith(".pdf"):
                extracted_text = extract_text_from_pdf(temp_path)
            elif temp_path.lower().endswith(".docx"):
                extracted_text = extract_text_from_docx(temp_path)
            else:
                messages.error(request, "Unsupported file format.")
                return redirect(f'../geteachDio/{id}')

            if not extracted_text:
                messages.error(request, "Could not extract text from the document.")
                return redirect(f'../geteachDio/{id}')

            extracted_text = clean_text(extracted_text)

            if not contains_required_words(extracted_text, FIXED_REQUIRED_WORDS, lecturer_names):
                messages.error(request, f"Document must contain all of these words: {', '.join(FIXED_REQUIRED_WORDS)} and at least one lecturer's name.")
                return redirect(f'../geteachDio/{id}')

            #  Check similarity
            existing_texts = [clean_text(doc.summary) for doc in Document.objects.all() if doc.summary]
            similarity_score = calculate_similarity(extracted_text, existing_texts)
            if similarity_score >= 0.35:
                theSimilarity = f"{similarity_score * 100:.2f}%"
                request.session['SimScore'] = theSimilarity
                messages.error(request, f"A similar project already exists with {theSimilarity} similarity.")
                return redirect(f'../geteachDio/{id}')

            #  Apply formatting before saving
            formatting = formatDocument.objects.first()
            if formatting:
                if temp_path.lower().endswith(".pdf"):
                    apply_format_to_pdf(temp_path, formatting)
                elif temp_path.lower().endswith(".docx"):
                    apply_format_to_docx(temp_path, formatting)

            #  Save the formatted file
            final_dir = os.path.join(settings.STATICFILES_DIRS[0], f'medias/{parentname}/{foldername}')
            os.makedirs(final_dir, exist_ok=True)
            final_path = os.path.join(final_dir, document.name)

            with open(temp_path, 'rb') as formatted_file:
                saved_name = default_storage.save(final_path, ContentFile(formatted_file.read()))

            file_url = default_storage.url(saved_name)
            file_uri = file_url.replace("/repository", "").replace("%20", " ")
            folder_urlx = file_uri.rsplit(foldername, 1)[0] + f'{foldername}/'
            status = "Not Approved"
            #  Save uploadFile and Document records
            uploadFile.objects.create(
                nameUploaded=nameUploaded,
                filename=filename,
                foldername=foldername,
                parentname=parentname,
                docs=file_uri,
                folderurl=folder_urlx,
                lecturer=lecturer,
                date_created=dateY,
                status=status
            )

            Mysummary = summarize_text(extracted_text)
            Document.objects.create(title=filename, file=file_uri, summary=Mysummary)

            messages.success(request, "Document uploaded and formatted successfully!")
            return redirect("document_list")

        except Exception as e:
            messages.error(request, f"An error occurred: {str(e)}")
            return redirect(f'../geteachDio/{id}')

    return redirect(f'../geteachDio/{id}')


# def geteachDio(request, id):
#     sim_score = request.session.pop('SimScore', None)
#     context = {
#         'simScore':sim_score
#     }
#     return render(request, 'geteachDio.html', context)


@login_required(login_url='Myhome')
def document_list(request):
    user_groups = request.user.groups.values_list('name', flat=True)
    if request.user.is_authenticated:
        firstname = request.user.first_name
    else:
        firstname = "General"
    
    documents = Document.objects.all()
    
    
    context = {"documents": documents, 'user_groups':user_groups, 'firstname':firstname}
    return render(request, "document_list.html",context)


def delete_files(request):
    print("Weldone")
    if request.method == "POST":
        file_ids = request.POST.getlist("file_ids")
        action = request.POST.get("action")  # Get selected action

        if file_ids and action == "delete":
            for file_id in file_ids:
                try:
                    file_obj = uploadFile.objects.get(id=file_id)
                    file_name = file_obj.filename
                    # Delete related documents by title
                    Document.objects.filter(title=file_name).delete()
                    # Delete the file itself
                    file_obj.delete()
                except uploadFile.DoesNotExist:
                    continue  # Skip if file not found

            messages.success(request, "Selected files deleted successfully.")
        else:
            messages.warning(request, "No action selected or no files checked.")
    
    return redirect("dashboard")


def signup(request):
    if request.method == "POST":
        username = request.POST["username"]
        email = request.POST["email"]
        password = request.POST["password"]

        if User.objects.filter(username=username).exists():
            messages.error(request, "Username already taken!")
        elif User.objects.filter(email=email).exists():
            messages.error(request, "Email already registered!")
        else:
            user = User.objects.create_user(username=username, email=email, password=password)
            user.save()
            student_group, created = Group.objects.get_or_create(name="Student")
            user.groups.add(student_group)
            messages.success(request, "Account created successfully! You can now login.")
            return redirect("Myhome")  # Redirect to login page

    return render(request, "signup.html")

@login_required(login_url='Myhome')
def documentFormating(request):
    FontSize = range(8,40)
    if request.method == "POST":
        font_size = request.POST.get("font_size")
        font_color = request.POST.get("font_color")
        letter_spacing = request.POST.get("letter_spacing")
        line_spacing = request.POST.get("line_spacing")
        font_family = request.POST.get("font_family")
        

        # Save to the database
        document, created = formatDocument.objects.update_or_create(
            id=1,  # You can adjust this if you want user-specific settings
            defaults={
                "fontSize": font_size,
                "fontColor": font_color,
                "letterSpacing": letter_spacing,
                "lineSpacing": line_spacing,
                "fontFamily": font_family,
                
            },
        )

        if created:
            messages.success(request, "Document formatting created successfully!")
        else:
            messages.success(request, "Document formatting updated successfully!")
        context = {'fontSize':FontSize}
        return render(request, 'documentFormating.html', context)

    # Fetch the existing document settings if available
    document = formatDocument.objects.first()

    context = {'fontSize':FontSize, 'document':document}
    return render(request, 'documentFormating.html', context)

@login_required(login_url='Myhome')
def showProject(request, id):
    user_groups = request.user.groups.values_list('name', flat=True)
    getNav = get_nav_context(request)
    getId =  uploadFile.objects.get(id=id)
    context = {'theFile':getId, **getNav, 'user_groups':user_groups,}
    if request.method == "POST":
        getID =  uploadFile.objects.get(id=id)
        if getID.status == "Approved":

            getID.status = "Not Approved"
            getID.save()
            return redirect(f'../showProject/{id}', context)
        else:
            getID.status = "Approved"
            getID.save()
            return redirect(f'../showProject/{id}', context)
        
    return render(request, 'editform.html', context)

def sendMessage(request, username):
    recipient = get_object_or_404(User, username=username)

    if request.method == "POST":
        message_text = request.POST.get("message")
        if message_text:
            Message.objects.create(
                sender=request.user,
                receiver=recipient,
                content=message_text
            )
        return redirect('sendMessage', username=recipient.username)

    # Fetch messages between sender and recipient
    messages = Message.objects.filter(
        Q(sender=request.user, receiver=recipient) |
        Q(sender=recipient, receiver=request.user)
    ).order_by('timestamp')

    # Group messages by date
    grouped_messages = defaultdict(list)
    for message in messages:
        date_key = localtime(message.timestamp).date()
        grouped_messages[date_key].append(message)

    context = {
        'recipient': recipient,
        'grouped_messages': grouped_messages.items(),  # date, [messages]
        'today': now().date(),
        'yesterday': (now() - timedelta(days=1)).date(),
    }

    return render(request, 'sendMessage.html', context)